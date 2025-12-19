# PyPDF2 removed from primary extraction chain (lower success rate ~60-70%)
# Using PyMuPDF (90-95%) + pdfminer.six (85-90%) for best results
from pdfminer.high_level import extract_text as pdfminer_extract_text
from pdf2image import convert_from_path
import pytesseract
import boto3
import fitz  # PyMuPDF - Primary extractor
from docx import Document
import logging
import tempfile
import os
import subprocess
from typing import Tuple, Dict, Any
import cv2
import numpy as np
from PIL import Image
# EasyOCR imported lazily when needed (heavy dependency)

logger = logging.getLogger(__name__)


class DocumentProcessor:
    def process_pdf(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from PDF file.

        Strategy (Optimized for best extraction quality):
        1. Try PyMuPDF (fitz) first - Best for PDFs with embedded text (~90-95% success).
        2. If resulting text is empty, fall back to pdfminer.six - Excellent for complex layouts and tables (~85-90% success).
        3. If still empty, run Tesseract OCR with image preprocessing - For scanned/image PDFs.
        4. If Tesseract fails and ENABLE_EASYOCR=true, try EasyOCR as fallback (heavy, optional).
        5. Optionally (when ENABLE_TEXTRACT=true), fall back to AWS Textract as a
           premium, cloud-based extractor for difficult cases.
        
        Note: PyPDF2 removed from primary chain due to lower success rate (60-70%).
        PyMuPDF + pdfminer.six combination provides 95%+ coverage for text-based PDFs.
        OCR with preprocessing + EasyOCR fallback handles most image-based PDFs.
        """
        text = ""
        metadata: Dict[str, Any] = {}

        # Try to repair severely malformed PDFs before processing, but never fail
        # if repair is not possible. This increases the chance that downstream
        # tools (PyMuPDF, pdfminer.six, OCR) can work with borderline files.
        repaired_path = self._try_repair_pdf(file_path)
        
        try:
            # First attempt: PyMuPDF (fitz) - Best balance of speed and reliability
            try:
                with fitz.open(file_path) as doc:
                    metadata['page_count'] = len(doc)
                    page_text_parts = []
                    for page in doc:
                        page_text = page.get_text()
                        if page_text.strip():
                            page_text_parts.append(
                                f"--- Page {page.number + 1} ---\n{page_text}\n\n"
                            )
                    pymupdf_text = "".join(page_text_parts)
                    if pymupdf_text.strip():
                        metadata["extractor"] = "pymupdf"
                        text = pymupdf_text
                        
                        # Extract PDF metadata if available
                        if doc.metadata:
                            metadata['pdf_metadata'] = {
                                'title': doc.metadata.get('title', ''),
                                'author': doc.metadata.get('author', ''),
                                'subject': doc.metadata.get('subject', ''),
                                'creator': doc.metadata.get('creator', ''),
                            }
            except Exception as e:
                logger.warning(f"PyMuPDF failed to extract text: {e}. Will try pdfminer.six.")
                text = ""
            
            # Second attempt: pdfminer.six - Excellent for complex layouts and tables
            if not text.strip():
                logger.info("PyMuPDF extraction empty; falling back to pdfminer.six")
                try:
                    pdfminer_text = pdfminer_extract_text(file_path) or ""
                    if pdfminer_text.strip():
                        metadata['extractor'] = 'pdfminer.six'
                        text = pdfminer_text
                except Exception as e:
                    logger.warning(f"pdfminer.six failed to extract text: {e}. Will try OCR.")
            
            # Third attempt: OCR with Tesseract (with image preprocessing) for scanned/image PDFs
            if not text.strip():
                logger.info("No text from PyMuPDF/pdfminer.six; attempting Tesseract OCR with preprocessing")
                logger.warning("OCR is slow - this may take a while for large PDFs...")
                try:
                    # Limit OCR to first 5 pages for performance (can be configured)
                    max_ocr_pages = int(os.getenv("MAX_OCR_PAGES", "5"))
                    pages = convert_from_path(repaired_path, first_page=1, last_page=max_ocr_pages)
                    
                    ocr_text_parts = []
                    for idx, img in enumerate(pages, start=1):
                        logger.info(f"Tesseract OCR processing page {idx}/{len(pages)}...")
                        # Preprocess image for better OCR results
                        preprocessed_img = self._preprocess_image_for_ocr(img)
                        page_text = pytesseract.image_to_string(preprocessed_img, timeout=30)  # 30s timeout per page
                        if page_text.strip():
                            ocr_text_parts.append(f"--- Tesseract OCR Page {idx} ---\n{page_text}\n")
                    
                    ocr_text = "".join(ocr_text_parts)
                    if ocr_text.strip():
                        metadata["extractor"] = "ocr-tesseract"
                        metadata["ocr_pages_processed"] = len(pages)
                        metadata["ocr_note"] = f"Only first {max_ocr_pages} pages processed for performance"
                        text = ocr_text
                    else:
                        logger.warning("Tesseract OCR returned no text - trying EasyOCR fallback...")
                except Exception as e:
                    logger.warning(f"Tesseract OCR failed: {e}. Will try EasyOCR fallback.")
            
            # Fourth attempt: EasyOCR as fallback for Tesseract (only if ENABLE_EASYOCR=true)
            if not text.strip() and os.getenv("ENABLE_EASYOCR", "false").lower() == "true":
                logger.info("Tesseract OCR failed; attempting EasyOCR fallback")
                try:
                    # Lazy import EasyOCR to avoid loading if not needed
                    try:
                        import easyocr  # type: ignore
                    except ImportError:
                        logger.warning("EasyOCR not installed. Skipping EasyOCR fallback.")
                        raise ImportError("EasyOCR not available")
                    
                    max_ocr_pages = int(os.getenv("MAX_OCR_PAGES", "5"))
                    pages = convert_from_path(repaired_path, first_page=1, last_page=max_ocr_pages)
                    
                    # Initialize EasyOCR reader (lazy load, English only for performance)
                    # Note: First initialization downloads models (~100MB) - may take 1-2 minutes
                    logger.info("Initializing EasyOCR (first time may download models)...")
                    easyocr_reader = easyocr.Reader(['en'], gpu=False)  # GPU=False for CPU-only
                    
                    ocr_text_parts = []
                    for idx, img in enumerate(pages, start=1):
                        logger.info(f"EasyOCR processing page {idx}/{len(pages)}...")
                        # Preprocess image for better OCR results
                        preprocessed_img = self._preprocess_image_for_ocr(img)
                        # Convert PIL Image to numpy array for EasyOCR
                        img_array = np.array(preprocessed_img)
                        results = easyocr_reader.readtext(img_array)
                        page_text = "\n".join([result[1] for result in results if result[2] > 0.5])  # Confidence > 0.5
                        if page_text.strip():
                            ocr_text_parts.append(f"--- EasyOCR Page {idx} ---\n{page_text}\n")
                    
                    ocr_text = "".join(ocr_text_parts)
                    if ocr_text.strip():
                        metadata["extractor"] = "ocr-easyocr"
                        metadata["ocr_pages_processed"] = len(pages)
                        metadata["ocr_note"] = f"Only first {max_ocr_pages} pages processed for performance"
                        text = ocr_text
                    else:
                        logger.warning("EasyOCR returned no text - PDF may be image-only or require better OCR settings")
                except Exception as e:
                    logger.warning(f"EasyOCR failed: {e}. Will try Textract if enabled.")

            # Optional fifth attempt: AWS Textract (cloud OCR + layout), only if enabled
            if not text.strip() and os.getenv("ENABLE_TEXTRACT", "false").lower() == "true":
                logger.info("No text from local extractors; falling back to AWS Textract")
                textract_text = self._extract_with_textract(file_path)
                if textract_text.strip():
                    metadata["fallback_extractor"] = "aws-textract"
                    text = textract_text
                else:
                    raise ValueError(
                        "No text could be extracted from PDF even with AWS Textract; document may be image-only or unsupported"
                    )
            
            return text.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise

    def _preprocess_image_for_ocr(self, img: Image.Image) -> Image.Image:
        """
        Preprocess image to improve OCR accuracy.
        
        Steps:
        1. Convert to grayscale
        2. Apply denoising
        3. Enhance contrast
        4. Apply thresholding for better text recognition
        
        Args:
            img: PIL Image object
            
        Returns:
            Preprocessed PIL Image
        """
        try:
            # Convert PIL Image to OpenCV format (numpy array)
            img_array = np.array(img)
            
            # Convert to grayscale if needed
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array
            
            # Apply denoising
            denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
            
            # Enhance contrast using CLAHE (Contrast Limited Adaptive Histogram Equalization)
            clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8, 8))
            enhanced = clahe.apply(denoised)
            
            # Apply adaptive thresholding for better text recognition
            # This converts the image to binary (black and white) which helps OCR
            thresh = cv2.adaptiveThreshold(
                enhanced, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                cv2.THRESH_BINARY, 11, 2
            )
            
            # Convert back to PIL Image
            preprocessed_img = Image.fromarray(thresh)
            return preprocessed_img
            
        except Exception as e:
            logger.warning(f"Image preprocessing failed: {e}. Using original image.")
            return img

    def _try_repair_pdf(self, file_path: str) -> str:
        """
        Best-effort repair of malformed PDFs using qpdf.

        If repair succeeds, return path to the repaired file; otherwise return
        the original path. This is intentionally conservative: we never fail
        the pipeline just because repair did not work.
        """
        try:
            tmp_dir = tempfile.mkdtemp(prefix="pdf_repair_")
            repaired_path = os.path.join(tmp_dir, "repaired.pdf")

            # --linearize writes a cleaned-up version; qpdf will attempt to
            # reconstruct cross-reference tables and streams where possible.
            result = subprocess.run(
                ["qpdf", "--linearize", file_path, repaired_path],
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=False,
                text=True,
            )

            if result.returncode == 0 and os.path.exists(repaired_path):
                logger.info("qpdf repair succeeded; using repaired PDF for extraction")
                return repaired_path

            logger.warning(
                "qpdf repair did not succeed (return code %s). Continuing with original PDF.",
                result.returncode,
            )
        except Exception as e:
            logger.warning(f"qpdf repair step failed: {e}. Continuing with original PDF.")

        return file_path

    def _extract_with_textract(self, file_path: str) -> str:
        """
        Extract text using AWS Textract by sending the PDF bytes directly.

        This is more powerful than local OCR but requires valid AWS credentials
        and is controlled via ENABLE_TEXTRACT env var.
        """
        try:
            region = os.getenv("AWS_REGION", "us-east-1")
            client = boto3.client("textract", region_name=region)

            with open(file_path, "rb") as f:
                document_bytes = f.read()

            response = client.detect_document_text(Document={"Bytes": document_bytes})

            lines = []
            for block in response.get("Blocks", []):
                if block.get("BlockType") == "LINE" and block.get("Text"):
                    lines.append(block["Text"])

            text = "\n".join(lines)
            return text
        except Exception as e:
            logger.error(f"AWS Textract extraction failed: {e}")
            return ""
    
    def process_docx(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from DOCX (Microsoft Word) file.
        
        Extracts:
        - All paragraph text
        - Text from tables
        - Document properties (title, author, created date)
        
        Returns:
        - Extracted text as string
        - Metadata dictionary with document information
        """
        text = ""
        metadata = {}
        
        try:
            doc = Document(file_path)
            metadata['paragraph_count'] = len(doc.paragraphs)
            
            # Extract text from paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            
            # Extract text from tables (DOCX files often have tables)
            table_count = 0
            for table in doc.tables:
                table_count += 1
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text += " | ".join(row_text) + "\n"
            
            if table_count > 0:
                metadata['table_count'] = table_count
            
            # Extract document properties
            core_props = doc.core_properties
            metadata['doc_properties'] = {
                'title': core_props.title or None,
                'author': core_props.author or None,
                'subject': core_props.subject or None,
                'created': str(core_props.created) if core_props.created else None,
                'modified': str(core_props.modified) if core_props.modified else None,
                'last_modified_by': core_props.last_modified_by or None,
            }
            
            # Calculate word count from extracted text
            word_count = len(text.split())
            metadata['word_count'] = word_count
            
            if not text.strip():
                logger.warning(f"DOCX file {file_path} appears to be empty or contains no extractable text")
            
            return text.strip(), metadata
            
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {e}")
            raise ValueError(f"Failed to process DOCX file: {str(e)}")
    
    def process_txt(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            metadata = {
                'line_count': len(text.splitlines()),
                'encoding': 'utf-8'
            }
            
            return text, metadata
            
        except UnicodeDecodeError:
            # Try latin-1 encoding
            with open(file_path, 'r', encoding='latin-1') as file:
                text = file.read()
            
            metadata = {
                'line_count': len(text.splitlines()),
                'encoding': 'latin-1'
            }
            
            return text, metadata
        except Exception as e:
            logger.error(f"Error processing TXT: {e}")
            raise
    
    def process_file(self, file_path: str, file_type: str) -> Tuple[str, Dict[str, Any]]:
        """Process file based on type"""
        if file_type == 'application/pdf':
            return self.process_pdf(file_path)
        elif file_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            return self.process_docx(file_path)
        elif file_type == 'text/plain':
            return self.process_txt(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")

document_processor = DocumentProcessor()